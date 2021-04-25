"""
Ворд:
1) Что если в таблице удалена ячейка в строке? В этом случае, нужно вывести ошибку про нарушение структуры таблицы.
2) Формулы. Библиотека питона их не видит. Для него формула -- пустота. Но если в ячейке есть обычный текст, то он
его отлично видит и отобразит его так же как в ячейке, но без формулы.
3) Рисунки. Тоже самое, что и с формулами.
3.1)Таблица внутри ячейки. Тоже что и формулы.
4) Слияние/объеденние ячеек работает без багов
5) Полужирный , курсив, нижнее подчеркивание, Нижний/вверхний индекс -- есть и работает стабильно.
6) Границы. Поддерживается только одна толщина и отображаются везде, кроме тех случаев, когда речь про
 объеденные ячейки.
7) Ширина ячейки. Тут есть проблемы. Латех не правильно отображает сантиметры , а ворд не очень корректно выдает
данные по ширине. К примеру, если в ячейке (в ворде) есть формула или же картинка, то ширина картинки/формулы
не учитывается. Видимо, ворд подставляет размер картинки во время внутренной работы, а в XML хранит размер без
учета формул/картинки.
Баг 7 пункта: Как оказалось, если ширина ячейки не указана (ворде можно такое сделать), то вызывается ошибка. Позже
исправить.
8) Выравнение. Если речь про то, что текст может быть слева/справа/по-центру, то проблем нету. Однако, если речь про
то что текст можеть быть сверху/внизу/центру, то тут уже проблемы. Как оказалось, метод, которым я это делаю,
имеет проблемы с том, что если в одной строке в одной ячейке  указано, что текст сверху, а в другой ячейке текст внизу,
то так не может быть. Каждая ячейка наследует поведение прошлой ячейки. Поэтому, если впервой ячейке текст сверху, то
все ячейки одной строки будут иметь текст с сверху. Кроме того, библиотека для работы с вордом не корректно отображает
выравнивание текст. Поэтому было решено, что весь текст в таблицах будет по центру. Но если он находиться справа или
слева, то это будет учтено. Если же сверху/внизу, то это не будет учтено и текст будет по середение высоты ячейки.
9) Распределение текста в ячейке. multirow пока-что сопротивляется, но думаю все будет с ним.
10) Фон/цвет ячеек. Как я понял это не нужно. А раз не нужно -- не делаеться.
11) Шрифт текста. Я когда-то изучал латех, то выявил, что латех не очень имеет дружбу с шрифтами ворда. В теории, я могу
вывести шрифт ворда, но сформировать его в латехе.... Сложно и как по мне не имеет смысла. Поэтому, таблицы наследуют
тот же шрифт, что используется в тех-документе.
12) Размер текста. В теории возможно, но в латехе странные функции по изменению размера.
13) В ворде можно поставить двойное зачеркивание текста. В латехе не нашел команды. Нужно ли?
14) Размер шрифта. Все было бы хорошо, если бы не одно НО. Если размер шрифта таблички совпадает с размером шрифта по
документу, то я не могу узнать его узнать -- функция выведет None, а не размер кегля. Узнать размер общий шрифт докумен-
та не очень-то и можно. Решил так. Если размер кегля совпадает с размером документа в ворде, то значит кегль таблицы
в латехе наследуется от общей настройки в  команде \documentclass[a4paper,12pt]{article}. В противном случае, размер
кегля таблицы будет корректироваться командой \small. Тоесть, если кегль документа 14pt, то табличка будет 12pt.
После синтаксиса таблицы, будет выводиться \normalsize.
Как я понял, кегль внутри таблички должен быть одинаковым. К тому же, если изменять размер каждого слова, то текст
начинает не очень адекватно выводиться и может вывестись весь в одну строку.
15) Табуляция в таблицах. Решено было пропустить.
16) Междустрочный интервал. Наследуется от латеха.
17) Специальные символы латеха обрабатываются

Что еще нужно сделать:
1) Специальная команда в начале tex-файл, что выводит необходимые команды в начале предложения. После begin{document}
должна оставаться. Также без этого не должны работать генератор. +
2) Подключение ворд-документа.Подключать можно где угодно, но команда должна занимать одну строку. +
3) Вывод ячеек в любой части латеха. Просто как текст. +
4) Разработка несколько режимов генератора:
-- Вывод таблицы как можно больше похоже на оригинал. Ха-ха. Original +
-- Вывод таблицы, где Текст без форматирования. Only_Text +
-- Вывод таблицы, где вместо текста команды для каждой ячейки. Only_Cells +
5) Внедрение генератора в основную таблицу. +
"""
# -----------------------------------------------------------
# Codes other files project
# -----------------------------------------------------------

import data.XML.work_with_XML as XML # Work with XML-file

# -----------------------------------------------------------
# Few libraries
# -----------------------------------------------------------

import os # work with os structure
import docx # work with structure
import re # work with regular expressions


class Generation_latex_word:
    # Search file-docx
    def search_docx_file(self, name_address=None, name_table=None, command_status = "Original"):
        # Checking that the user has not provided data
        if name_address == None:
            return ("Не вказана назва файлу!")
        if name_table == None:
            return "Не вказана назва таблиці!"
        # To make it easier to work in the class
        self.name_address = name_address
        self.name_table = name_table
        # Checking that the file exists at all
        if os.path.isfile(self.name_address):
            self.doc = docx.Document(self.name_address)
            self.properties = self.doc.core_properties  # paramaters docx-file
            return self.search_docx_tables(command_status)  # Run table search
        else:  # In case the file was not found.
            return "Файл не знайдено"

    # Search docx tables
    def search_docx_tables(self,command_status):
        i = 0  # Specifies the number of the required table.
        find_table = False  # Table location marker
        for tables in self.doc.tables:  # Find the table
            if find_table == True:  # Found a table!
                    break
            if tables.rows[0].cells[0].text == self.name_table:  # If found a table
                return self.Assembly_shop(i,command_status)
            else:
                i += 1
        if find_table == False:  # If the table was not found
            return "Таблиця не знайдена"

    # Determines the width and position of the text in the cell
    def align_multirow(self, row, column):
        # Defines the width of the cell in cm
        try:
            a = (self.name_table.cell(row, column).width / 914400.0) * 2.54
            a = round(a*0.9,1)
        # Default width of the cell in cm -- 1 cm
        except Exception:
            a = 1.0
        # Dress width cell into parbox tag
        string = "\\parbox{" + str(a) + "cm" + "}{"
        a_aligment = str(self.name_table.cell(row, column).paragraphs[0].alignment)
        # If the cell text is on the right side
        if a_aligment == "RIGHT (2)":
            string += " \\raggedleft "
            return string
        # If the cell text is in the center of the cell
        elif a_aligment == "CENTER (1)":
            string += " \\centering "
            return string
        # If the cell text is on the left side
        elif a_aligment == "LEFT (4)":
            string += " \\raggedright "
            return string
        # If the cell text fits the width of the cell
        elif a_aligment == "JUSTIFY (3)":
            return string
        # If the positioning of the text is not determined,
        # then the text is automatically centered
        else:
            return string

    # Processing cells in the table
    def every_row(self,command_status):
        # Drawing up a map of occupied cells.
        # This is necessary so that cells are not processed
        # Which have already merged with the cell that is defined
        # as a cell with merge
        # 0 - the cell is not occupied
        # 1 - the cell is occupied
        map_table = [[0] * len(self.name_table.columns) for i in range(len(self.name_table.rows))]
        # Processing each row in the table
        for row in range(len(self.name_table.rows)):
            # row 0 is row where only name-tables
            if row == 0:
                pass
            else:
                string = ''
                column = 0
                # Same as "map of occupied cells"
                # Read the description in the same place
                # Needed for border between lines
                map_table_hline = [""] * len(self.name_table.columns)
                # Processing each column in the table
                while column < len(self.name_table.columns):
                    f = 1  # Counter for concatenating cells by row
                    c = 0  # Counter for concatenating cells by column

                    # Defines the border for the steering column
                    if column == 0:
                        border_left = "|"
                    else:
                        border_left = ""
                    border_right = ""  # Variable to define the border between the columns

                    # # If the cell is not merged with someone
                    if map_table[row][column] == 0:
                        k = 0
                        # print(str(row + 1) + " " + str(column+1))
                        #self.merge_cells_word(row,column) -- МОЖЕТ БЫТЬ ЭТО НАДО БУДЕТ ВКЛЮЧИТЬ в СЛУЧАЕ ПОЛОМКИ

                        # If the merge happened both by columns and by rows
                        if row < (len(self.name_table.rows) - 1) and column < (len(self.name_table.columns) - 1)\
                                and self.merge_cells_word(row,column) == self.merge_cells_word(row + 1, column)\
                                and self.merge_cells_word(row,column) == self.merge_cells_word(row, column+1):

                            #----------------------

                            # The loop checks which cells belong to
                            # the cell in which the merge was detected
                            while (k + column) < (len(self.name_table.columns) - 1):
                                # If found by column
                                if self.merge_cells_word(row, k + column) == self.merge_cells_word(row,k + column + 1):
                                    f += 1
                                else:
                                    break
                                k += 1
                            # Change border status
                            border_right = "|"
                            map_table_hline[column+f-1] = "|"
                            string += "\\multicolumn{" + str(f) +"}{" + border_left + self.paragraphs_alignment_cell(row,column) + border_right +"}{"
                            c += 1
                            row_merge = 1
                            # The loop checks which cells belong to
                            # the cell in which the merge was detected
                            while (row_merge + row) < len(self.name_table.rows):
                                # If found by rows
                                if self.merge_cells_word(row, column) == self.merge_cells_word(row_merge + row, column):
                                    k = 0
                                    while (k + column) < (column+f-1):
                                        map_table[row_merge + row][k + column] = 2
                                        k += 1
                                    map_table[row_merge + row][k + column] = 2 # Last column to merge_cell
                                    c += 1
                                else:
                                    break
                                row_merge += 1
                            # If the end of the columns in the line
                            if column+f == len(self.name_table.columns):
                                string += "\multirow{" + str(c) + "}{*}{" + self.align_multirow(row, column)+ self.return_cells(row, column,command_status) + "}}}"
                            else:
                                string += "\multirow{" + str(c) + "}{*}{" + self.align_multirow(row, column) + self.return_cells(row, column,command_status) + "}}}" + "&"

                            #-----------------------

                        # If combining only by  rows
                        elif row < (len(self.name_table.rows) - 1) \
                                and self.merge_cells_word(row,column) == \
                                self.merge_cells_word(row + 1, column):
                            map_table[row][column] = 1
                            map_table[row + 1][column] = 1
                            c += 1
                            row_merge = 1
                            border_right = "|"
                            map_table_hline[column + f-1] = "|"
                            string += "\\multicolumn{" + str(1) + "}{" + border_left + self.paragraphs_alignment_cell(row,column) + border_right + "}{"
                            # Merge by rows
                            while row_merge + row < len(self.name_table.rows):
                                if self.merge_cells_word(row, column) == self.merge_cells_word(row_merge + row, column):
                                    map_table[row_merge + row][column] = 1
                                    c += 1
                                else:
                                    break
                                row_merge += 1
                                map_table_hline[column+f-1] = "|"
                            # If the last column
                            if (column + f) == len(self.name_table.columns):
                                string += "\\multirow{" + str(c) + "}{*}{" + self.align_multirow(row, column)+ \
                                          self.return_cells(row,column,command_status) + "}}}"
                            # If the first column
                            elif column == 0:
                                string += "\\multirow{" + str(c) + "}{*}{" + self.align_multirow(row, column)+ \
                                          self.return_cells(row,column,command_status) + "}}}" + "&"
                            # If in other cases
                            else:
                                string += "\\multirow{" + str(c) + "}{*}{" + self.align_multirow(row, column)+ self.return_cells(row,column,command_status) + "}}}" + "&"
                        # Merge by columns
                        else:
                            while (k + column) < (len(self.name_table.columns) - 1):
                                if self.merge_cells_word(row, k + column) == self.merge_cells_word(row,
                                                                                                   k + column + 1):
                                    f += 1
                                else:
                                    break
                                k += 1
                            #|-----------|
                            #|           |
                            #|-----------|
                            if column + f == len(self.name_table.columns): # Last cell in row
                                ampersand = ""
                            else:
                                ampersand = "&"
                            border_right = "|"
                            map_table_hline[column + f-1] = "|"
                            string += "\\multicolumn{" + str(f) + "}{"+ border_left +self.paragraphs_alignment_cell(row,column) + border_right+"}{" + \
                                          self.return_cells(row,column,command_status) + "}" + ampersand
                    # # If the cell was previously merged
                    else:
                        # map-table == 1... or 0
                        # Last cell in row
                        if column + f == len(self.name_table.columns):
                            ampersand = ""
                        else:
                            ampersand = "&"
                        # Last cell in row
                        if column + f == len(self.name_table.columns):
                            border_right = "|"
                            map_table_hline[column + f-1] = "|"
                        # If the cell is not the last in the row and does not merge with the neighbor to the right
                        elif column + f != len(self.name_table.columns) and self.merge_cells_word(row,column) != self.merge_cells_word(row,column+1):
                            border_right = "|"
                            map_table_hline[column + f-1] = "|"
                        if map_table == 1:
                            string += "\\multicolumn{" + str(f) + "}{" + border_left + self.paragraphs_alignment_cell(row,column) + border_right + "}{" + "}" + ampersand
                        else:
                            string += "\\multicolumn{" + str(f) + "}{" + border_left + "c" + border_right + "}{" + "}" + ampersand
                    # Skip the pooled cells
                    column = column + f
                self.list_new_table.append(string + "\\\\")
                # Create a border between the line on the border map
                hhline = "\\hhline{"
                for column_hline in range(len(self.name_table.columns)):
                    if row == (len(self.name_table.rows) - 1):
                        hhline += "-"
                    elif map_table[row+1][column_hline] == 0:
                        hhline += "-"
                    elif column_hline == 0:
                        hhline += "|~" + map_table_hline[column_hline]
                    else:
                        hhline += "~" + map_table_hline[column_hline]
                hhline += "}"
                # Добавляем ячейки в таблицу
                self.list_new_table.append(hhline)


class Generation_latex(Generation_latex_word):

    def __init__(self):
        self.doc = ""
        self.name_address = ""
        self.name_table = ""
        self.list_start_file = list()
        self.list_new_table = list()
        self.list_docx_file = list()

    # Determines with whom the cell is merged
    # For this, an ID of the cell is compiled
    # It consists of the relationship between the cell and
    # of its neighboring cells
    # ID of the pooled cells is the same
    def merge_cells_word(self, row, column):
        try:
            return str(self.name_table.cell(row, column)._tc.top) + str(self.name_table.cell(row, column)._tc.bottom) + \
                   str(self.name_table.cell(row, column)._tc.left) + str(self.name_table.cell(row, column)._tc.right)
        except Exception:
            # In case the merged cells occur between the lines
            return str(self.name_table.cell(row, column)._tc.top) + "0707" + \
               str(self.name_table.cell(row, column)._tc.left) + str(self.name_table.cell(row, column)._tc.right)

    def paragraphs_alignment_cell(self,row,column):
        # Defines the width of the cell in cm
        try:
            a = (self.name_table.cell(row, column).width / 914400.0) * 2.54
            a = round(a*0.9,1)
        # Default width of the cell in cm -- 1 cm
        except Exception:
            a = 1.0
        a_aligment = str(self.name_table.cell(row, column).paragraphs[0].alignment)
        # If the cell text is on the right side
        if a_aligment == "RIGHT (2)":
            return (">{\\raggedleft\\arraybackslash}m{" + str(a) + "cm" + "}")
        # If the cell text is on the center side
        elif a_aligment == "CENTER (1)":
            return (">{\\centering\\arraybackslash}m{" + str(a) + "cm" + "}")
        # If the cell text is on the left side
        elif a_aligment == "LEFT (4)":
            return(">{\\raggedright\\arraybackslash}m{" + str(a) + "cm" + "}")
        # If the cell text is on the JUSTIFY side
        elif a_aligment == "JUSTIFY (3)":
            return ("m{" + str(a) + "cm" + "}")
        # If the positioning of the text is not determined,
        # then the text is automatically right
        else:
            return (">{\\raggedright\\arraybackslash}m{" + str(a) + "cm" + "}")

    # Specifies that the text is bold
    def run_bold(self,string_r,status):
        if status == "True":
            return "\\textbf{" + string_r +"}"
        else:
            return string_r

    # Specifies that the text is italic
    def run_italic(self,string_r,status):
        if status == "True":
            return "\\textit{" + string_r +"}"
        else:
            return string_r

    # Specifies that the text is underline
    def run_underline(self,string_r,status):
        if status == "True":
            return "\\underline{" + string_r +"}"
        else:
            return string_r

    # Specifies that the text is strikethrough
    def run_font_strike(self,string_r,status):
        if status == "True":
            return "\\sout{" + string_r + "}"
        else:
            return string_r

    # Specifies that the text is subscript
    def run_font_subscript(self,string_r,status):
        if status == "True":
            return "True"
        else:
            return "False"

    # Specifies that the text is superscript
    def run_font_superscript(self,string_r,status):
        if status == "True":
            return "True"
        else:
            return "False"

    # Specifies that the text is superscript
    # If the text size of the word table is different
    # from the size of the word document
    # then the size of the table text in latex is set a little less
    # otherwise the size of the table text is inherited from latex
    def table_font_size(self):
        try:
            status_size = str(self.name_table.cell(1, 0).paragraphs[0].runs[0].font.size)
        except Exception:
            status_size = "None"
        if status_size == "None":
            return "False"
        else:
            return "True"

    # Handling latex special characters
    def special_symbols(self, string):
        # List of latech symbols
        mas_a = ['\\','#','$','%','^', '&', '_', '{","}','~']
        # Latech character processing list
        mas_b = ['\\textbackslash ','\\# ','\\$ ','\\% ','\\textasciicircum ', '\\& ', '\\_ ', '\\{ ","\\} ','\\textasciitilde ']
        for i in range(len(mas_a)):
            string = string.replace(mas_a[i],mas_b[i])
        return string

    def return_cells(self,row,column,command_status):
        # Width of each cell
        string_p = ""
        string_r = ""
        # Cell text processing loop
        for paragraph in self.name_table.cell(row, column).paragraphs:
            string_r = ""
            # Loop processing each piece of cell text
            for run in paragraph.runs:
                string_rr = run.text
                string_rr = self.special_symbols(string_rr)
                math_first = ""
                math_formula = ""
                math_last = ""
                # Specifies that the text is subscript
                if self.run_font_subscript(string_rr, str(run.font.subscript)) == "True":
                    math_first = "$"
                    math_formula = "_{"
                    math_last = "}$"
                    # Specifies that the text is superscript
                if self.run_font_subscript(string_rr, str(run.font.superscript)) == "True":
                    math_first = "$"
                    math_formula = "^{"
                    math_last = "}$"

                if string_rr != (" " or ". ") and command_status == "Original":
                    # Specifies that the text is bold
                    string_rr = self.run_bold(string_rr, str(run.bold))
                    # Specifies that the text is italic
                    string_rr = self.run_italic(string_rr, str(run.italic))
                if command_status == "Original":
                    # Specifies that the text is underline
                    string_rr = self.run_underline(string_rr, str(run.underline))
                    # Specifies that the text is strike
                    string_rr = self.run_font_strike(string_rr, str(run.font.strike))
                # Assembling a piece of cell text
                string_r += math_first + math_formula + string_rr + math_last
            string_p += string_r
            # line break
            if command_status == "Original":
                string_p += "\\pol "
        # removal of extra processing chunk
        if command_status == "Original":
            string_p = string_p[0:-5]
        string_p += ""
        # Two modes when text is returned with or without processing
        if command_status == "Original" or command_status == "Only_Text":
            return string_p
        # Mode when only cell numbers are returned
        elif command_status == "Only_Cells":
            a = " Cell(" + self.name_table.cell(0, 0).text + "," + str(row) + "," + str(column) + ") "
            return a
        else:
            return  self.name_table.cell(row, column).text

    # Beginning of table
    def commands_to_generation(self, status):
        if status == "True":
            self.list_start_file.append("\\usepackage{multirow} % Об'єднання через рядки")
            self.list_start_file.append("\\usepackage{hhline} % Міжрядкова лінія")
            self.list_start_file.append("\\usepackage{array} % Вирівнювання")
            self.list_start_file.append("\\usepackage[normalem]{ulem} % Закреслення тексту")
            self.list_start_file.append("\\newcommand{\pol} % Для переносів")
            self.list_start_file.append("{")
            self.list_start_file.append(" ")
            self.list_start_file.append(" ")
            self.list_start_file.append("}")
        else:
            self.list_start_file.append("%CommandsGenerationlatexpython")

    # Beginning of table
    def first_part_table(self):  #
        if self.table_font_size() == "True": self.list_new_table.append("\\small")
        self.list_new_table.append("\\begin{tabular}" + "{" + "c" * len(self.name_table.columns) + "}")
        self.list_new_table.append("    \\hline")

    # Build the table
    def Assembly_shop(self, number_table, command_status):
        # Catch the desired table
        self.name_table = self.doc.tables[number_table]
        try:
            # Check the table for integrity
            for row in range(len(self.name_table.rows)):
                for column in range(len(self.name_table.columns)):
                    self.merge_cells_word(row, column)
        except (IndexError, AttributeError):
            return "Таблиця має не правильну структуру"
        # Beginning of table
        self.first_part_table()
        # Processing table cells
        self.every_row(command_status)
        # End of table
        self.list_new_table.append("\\end{tabular}")
        # Return the size of the latex document to its normal state
        if self.table_font_size() == "True":
            self.list_new_table.append("\\normalsize")
        return self.list_new_table

    # breaks the Generationlatexpython command into parameters
    def find_parameter_to_command_to_latex_file(self, structure_command):
        status_parameters = "False"
        # List of table modes. Insert_Original -- Excessive
        # Original
        # The output of the table is as similar to the original as possible.
        # Only_Text
        # Output the table, where Text without formatting.
        # Only_Cells
        # Output a table, where instead of the command text for each cell.
        names_parameters = ["Original", "Insert_Original", "Only_Text", "Only_Cells"]
        # 1 parameter. Word address
        name_address = re.search(r'(?<=\{)([\s\S]+?)(?=\,)', structure_command)
        # 2 parameter. Table name
        name_table = re.search(r'(?<=\,)([\s\S]+?)(?=\})', structure_command)
        # Determines the format of the current document
        name_address_data = re.search(r'[^.]+$', name_address.group(0))
        # Defines table mode
        for name in names_parameters:
            name = r'(?<=\{)(' + name + r')(?=\})'
            try:
                command_status = re.search(name, structure_command)
                command_status = command_status.group(0)
                break
            # If not decided, then put Original
            except AttributeError:
                command_status = "Original"
        # We expose the command, but disabled
        self.list_new_table.append('%Generationlatexpython_Disable{'+name_address.group(0) + "," + name_table.group(0) + "}{" + command_status + "}")
        if name_address_data.group(0) == "docx":
            return self.search_docx_file(name_address.group(0), name_table.group(0),command_status)
        else:
            return "Формат документа може бути тільки docx"

    # Processing the Include Docx command
    def find_parameter_to_command_to_IncludeDocx(self, structure_command):
        # Find the command Include Docx
        name_address = re.search(r'(?<=\{)([\s\S]+?)(?=\})', structure_command)
        # If file exists
        if os.path.isfile(name_address.group(0)):
            name_address_data = re.search(r'[^.]+$', name_address.group(0))
            # If the file is format docx, then add to the list
            if name_address_data.group(0) == "docx":
                self.list_docx_file.append(name_address.group(0))
                return structure_command
            else:
                return structure_command + " Формат документа може бути тільки docx"  # Придумать чет по-лучше.
        elif name_address.group(0) == "Назва_файлу":
            return structure_command
        else:
            return structure_command + " Такого файлу не існує"

    # Processing the Cell command
    def return_many_cells(self, structure_command):
        my = "True"
        table_string = ""
        try:
            # Find the command Cell
            string = re.search(r'Cell\s*\([^(),]+,\d+,\d+\)', structure_command).group(0)
        except AttributeError:
            return structure_command + "Attribute_error"
        # Processing each Cell
        while my == "True":
            #print(string)
            name_table = re.search(r'(?<=\()([\s\S]+?)(?=,)', string).group(0)
            #print(name_table)
            row = re.search(r'(?<=,)([\s\S]+?)(?=,)', string).group(0)
            #print(row)
            column = re.search(r'(?<=,)([\s\S]+?)(?=\))', string).group(0)
            column = re.search(r'(?<=,)([\s\S]+?)', column).group(0)
            #print(row + "|" + column)
            #print(column)
            a = "Cell_disable(" + name_table + "," + row + "," + column + ") - "
            table_string = a + " Впишіть назву документа за допомогою команди %IncludeDocx{Назва_файлу}"
            for docx_file in self.list_docx_file:
                #Find the file
                if os.path.isfile(docx_file):
                    doc = docx.Document(docx_file)
                else:
                    table_string = a + " Впишіть назву документа за допомогою команди %IncludeDocx{Назва_файлу}"
                #Find the file
                i = 0  # Specifies the number of the required table.
                find_table = False  # Table location marker
                for tables in doc.tables:  # Find the table
                    if find_table == True:  # Found a table!
                        break
                    if (tables.cell(0,0).text == name_table):  # Are looking for
                        find_table = True
                        # Text processing by spec. Latex symbols
                        if row.isdigit() and column.isdigit() and isinstance(row,float) == False and isinstance(column,float)== False:
                            if int(row) < len(tables.rows) and int(row) >= 0 and int(column) < len(
                                    tables.columns) and int(column) >= 0:
                                table_string = self.special_symbols(tables.cell(int(row), int(column)).text)
                            else:
                                table_string = a + "Координати комірки вказані не вірно"
                        else:
                            table_string = a + "Координати комірки може бути тількі цілочисленними цифрами та вище нуля "
                    else:
                        i += 1
                # If the table was not found
                if find_table == False:
                    table_string = a + "Таблиця ( " + name_table+ ")не знайдена по адресу (" + str (docx_file) + ") не знайдена"
                    table_string = self.special_symbols(table_string)
            # Change the command to cell text
            structure_command = structure_command.replace(string,table_string)
            # Determine if there are more cells
            if str(re.search(r'Cell\s*\([^(),]+,\d+,\d+\)',structure_command)) == "None":
                my = "False"
            else:
                string = re.search(r'Cell\s*\([^(),]+,\d+,\d+\)', structure_command).group(0)
        return structure_command

    # Search for commands in the latex file
    def find_command_to_latex_file(self, name_address_tex_file,status_run):
        find_begin_document = "True"
        status_generation = "False"
        if os.path.isfile(name_address_tex_file):
            with open(name_address_tex_file, 'r+', encoding='utf-8') as file:
                for line in file:
                    # Marker that the main document started.
                    if re.search(r'\\begin{document}',line):
                        find_begin_document = "False"
                        line = re.sub("^\s+|\n|\r|\s+$", '', line)
                        self.list_start_file.append(line)
                    # Command Cell detection
                    elif re.search(r'Cell\s*\([^(),]+,\d+,\d+\)',line):
                        #print(line)
                        line = self.return_many_cells(line)
                        line = re.sub("^\s+|\n|\r|\s+$", '', line)
                        self.list_start_file.append(line)
                    # # Command IncludeDocx detection
                    elif re.search(r'%IncludeDocx\s*{([\s\S]+?)}', line):
                        line = re.sub("^\s+|\n|\r|\s+$", '', line)
                        self.list_start_file.append(self.find_parameter_to_command_to_IncludeDocx(line))
                    # Command CommandsGenerationlatexpython detection
                    elif re.search(r'%CommandsGenerationlatexpython',line):  # Добавление команд для генератора.
                        self.commands_to_generation(find_begin_document)
                        find_begin_document = "False"
                        status_generation = "True"
                    # Command Generationlatexpython detection
                    elif (re.search(r'%Generationlatexpython\s*{[^(),]+,[^(),]+}{[^(),]+}', line))\
                            and status_generation == "True":  # Находим непосредственно команду для генерации таблиц.
                        line = re.sub("^\s+|\n|\r|\s+$", '', line)
                        self.list_start_file.extend(self.find_parameter_to_command_to_latex_file(line))
                        self.list_new_table.clear()
                    # If it is a line with no special commands
                    else:
                        line = re.sub("^\s+|\n|\r|\s+$", '', line)
                        self.list_start_file.append(line)
            # For test
            if (status_run == "enable"):
                file_name = XML.get_osnova_XML('tec-address') + "/" + XML.get_osnova_XML('tec-name-file') + ".tex"
            elif status_run == "test":
                file_name = "test_table.tex"
            else:
                file_name = XML.get_osnova_XML('tec-address') + "/" + XML.get_osnova_XML('tec-name-file') + '_enable' + ".tex"
            MyFile = open(file_name, 'w',  encoding='utf-8')
            self.list_start_file = map(lambda x: x + '\n', self.list_start_file)
            MyFile.writelines(self.list_start_file)
            MyFile.close()
        else:
            # Pass?
            print("Файлу нема!")


if __name__ == '__main__':
    app = Generation_latex()
    app.find_command_to_latex_file('test.tex', "test")